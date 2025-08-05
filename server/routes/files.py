from flask import Blueprint, request, jsonify
from werkzeug.utils import secure_filename
import os
from datetime import datetime
import pandas as pd
import fitz  # PyMuPDF for PDF handling
from docx import Document
from PIL import Image
import io

files_bp = Blueprint('files', __name__)

UPLOAD_FOLDER = os.path.join(os.path.dirname(__file__), '..', 'uploads')
ALLOWED_EXTENSIONS = {
    'doc', 'docx',  # Word documents
    'pdf',          # PDF files
    'xls', 'xlsx',  # Excel files
    'png', 'jpg', 'jpeg', 'gif', 'webp',  # Images
    'txt'           # Text files
}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def process_excel(file_path):
    """Extract data from Excel files"""
    try:
        df = pd.read_excel(file_path)
        return {
            'headers': df.columns.tolist(),
            'rows': df.values.tolist(),
            'shape': df.shape
        }
    except Exception as e:
        return {'error': f'Failed to process Excel file: {str(e)}'}

def process_word(file_path):
    """Extract text from Word documents"""
    try:
        doc = Document(file_path)
        return {
            'paragraphs': [p.text for p in doc.paragraphs],
            'tables': [[[cell.text for cell in row.cells] for row in table.rows] 
                      for table in doc.tables]
        }
    except Exception as e:
        return {'error': f'Failed to process Word document: {str(e)}'}

def process_text(file_path):
    """Extract text from plain text files"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return {
            'content': content,
            'lines': content.split('\n'),
            'line_count': len(content.split('\n')),
            'character_count': len(content)
        }
    except Exception as e:
        return {'error': f'Failed to process text file: {str(e)}'}

def process_pdf(file_path):
    """Extract text and metadata from PDF files"""
    try:
        doc = fitz.open(file_path)
        content = []
        for page in doc:
            content.append({
                'page_number': page.number + 1,
                'text': page.get_text(),
                'images': len(page.get_images())
            })
        return {
            'pages': content,
            'metadata': doc.metadata
        }
    except Exception as e:
        return {'error': f'Failed to process PDF: {str(e)}'}

def process_image(file_path):
    """Process and analyze image files"""
    try:
        with Image.open(file_path) as img:
            return {
                'format': img.format,
                'mode': img.mode,
                'size': img.size,
                'info': img.info
            }
    except Exception as e:
        return {'error': f'Failed to process image: {str(e)}'}

@files_bp.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
        
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
        
    if not allowed_file(file.filename):
        return jsonify({'error': f'File type not allowed. Allowed types: {", ".join(ALLOWED_EXTENSIONS)}'}), 400

    try:
        # Secure the filename and generate unique name
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_filename = f'{timestamp}_{filename}'
        file_path = os.path.join(UPLOAD_FOLDER, unique_filename)
        
        # Save the file
        file.save(file_path)
        
        # Process file based on type
        file_type = filename.rsplit('.', 1)[1].lower()
        
        if file_type in ['xls', 'xlsx']:
            result = process_excel(file_path)
        elif file_type in ['doc', 'docx']:
            result = process_word(file_path)
        elif file_type == 'pdf':
            result = process_pdf(file_path)
        elif file_type in ['png', 'jpg', 'jpeg', 'gif', 'webp']:
            result = process_image(file_path)
        elif file_type == 'txt':
            result = process_text(file_path)
        else:
            result = {'error': 'Unsupported file type'}
            
        # Add file info to result
        result.update({
            'filename': unique_filename,
            'original_filename': filename,
            'file_type': file_type,
            'file_size': os.path.getsize(file_path)
        })
        
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'error': f'Failed to process file: {str(e)}'}), 500

@files_bp.route('/upload-batch', methods=['POST'])
def upload_batch():
    """Upload and process multiple files"""
    if 'files' not in request.files:
        return jsonify({'error': 'No files provided'}), 400
    
    files = request.files.getlist('files')
    
    if not files or len(files) == 0:
        return jsonify({'error': 'No files selected'}), 400
    
    results = []
    errors = []
    
    for file in files:
        try:
            if file.filename == '':
                errors.append({'error': 'Empty filename'})
                continue
                
            if not allowed_file(file.filename):
                errors.append({
                    'filename': file.filename,
                    'error': 'File type not allowed'
                })
                continue
            
            # Process each file
            filename = secure_filename(file.filename)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            unique_filename = f'{timestamp}_{filename}'
            file_path = os.path.join(UPLOAD_FOLDER, unique_filename)
            
            # Save the file
            file.save(file_path)
            
            # Process file based on type
            file_type = filename.rsplit('.', 1)[1].lower()
            
            if file_type in ['xls', 'xlsx']:
                result = process_excel(file_path)
            elif file_type in ['doc', 'docx']:
                result = process_word(file_path)
            elif file_type == 'pdf':
                result = process_pdf(file_path)
            elif file_type in ['png', 'jpg', 'jpeg', 'gif', 'webp']:
                result = process_image(file_path)
            elif file_type == 'txt':
                result = process_text(file_path)
            else:
                result = {'error': 'Unsupported file type'}
            
            # Add file info to result
            result.update({
                'filename': unique_filename,
                'original_filename': filename,
                'file_type': file_type,
                'file_size': os.path.getsize(file_path)
            })
            
            results.append(result)
            
        except Exception as e:
            errors.append({
                'filename': file.filename if file.filename else 'unknown',
                'error': f'Failed to process file: {str(e)}'
            })
    
    return jsonify({
        'results': results,
        'errors': errors,
        'total_files': len(files),
        'successful': len(results),
        'failed': len(errors)
    })

@files_bp.route('/files/<filename>', methods=['GET'])
def get_file(filename):
    """Retrieve processed file data"""
    file_path = os.path.join(UPLOAD_FOLDER, filename)
    
    if not os.path.exists(file_path):
        return jsonify({'error': 'File not found'}), 404
        
    try:
        file_type = filename.rsplit('.', 1)[1].lower()
        
        if file_type in ['xls', 'xlsx']:
            result = process_excel(file_path)
        elif file_type in ['doc', 'docx']:
            result = process_word(file_path)
        elif file_type == 'pdf':
            result = process_pdf(file_path)
        elif file_type in ['png', 'jpg', 'jpeg']:
            result = process_image(file_path)
        else:
            result = {'error': 'Unsupported file type'}
            
        return jsonify(result)
        
    except Exception as e:
        return jsonify({'error': f'Failed to process file: {str(e)}'}), 500
