from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict, Any, List
import os
import io
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from PIL import Image, ImageDraw, ImageFont
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from io import BytesIO
import base64

class ReportGenerator:
    def __init__(self):
        self.document = Document()
        self.styles = getSampleStyleSheet()
        
    def generate_report_all_formats(self, 
                                  company_info: Dict[str, Any],
                                  valuation_data: Dict[str, Any],
                                  market_data: Dict[str, Any],
                                  peer_comparison: List[Dict[str, Any]],
                                  output_dir: str = "reports") -> Dict[str, str]:
        """Generate reports in all formats and return file paths"""
        
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        company_name = company_info.get("name", "Company").replace(" ", "_")
        
        # Generate all format reports
        formats = {}
        
        # Word format
        word_path = os.path.join(output_dir, f"{company_name}_valuation_report_{timestamp}.docx")
        self.generate_word_report(company_info, valuation_data, market_data, peer_comparison, word_path)
        formats['docx'] = word_path
        
        # PDF format
        pdf_path = os.path.join(output_dir, f"{company_name}_valuation_report_{timestamp}.pdf")
        self.generate_pdf_report(company_info, valuation_data, market_data, peer_comparison, pdf_path)
        formats['pdf'] = pdf_path
        
        # Text format
        txt_path = os.path.join(output_dir, f"{company_name}_valuation_report_{timestamp}.txt")
        self.generate_text_report(company_info, valuation_data, market_data, peer_comparison, txt_path)
        formats['txt'] = txt_path
        
        # Image format
        img_path = os.path.join(output_dir, f"{company_name}_valuation_report_{timestamp}.png")
        self.generate_image_report(company_info, valuation_data, market_data, peer_comparison, img_path)
        formats['png'] = img_path
        
        return formats

    def generate_comprehensive_report_all_formats(self, 
                                                company_info: Dict[str, Any],
                                                valuation_data: Dict[str, Any],
                                                market_data: Dict[str, Any],
                                                peer_comparison: List[Dict[str, Any]],
                                                output_dir: str = "reports") -> Dict[str, str]:
        """Generate comprehensive valuation reports with all three methods in all formats"""
        
        # Create output directory if it doesn't exist
        os.makedirs(output_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        company_name = company_info.get("name", "Company").replace(" ", "_")
        
        # Generate all format reports
        formats = {}
        
        # Word format
        word_path = os.path.join(output_dir, f"{company_name}_comprehensive_valuation_{timestamp}.docx")
        self.generate_comprehensive_word_report(company_info, valuation_data, market_data, peer_comparison, word_path)
        formats['docx'] = word_path
        
        # PDF format
        pdf_path = os.path.join(output_dir, f"{company_name}_comprehensive_valuation_{timestamp}.pdf")
        self.generate_comprehensive_pdf_report(company_info, valuation_data, market_data, peer_comparison, pdf_path)
        formats['pdf'] = pdf_path
        
        # Text format
        txt_path = os.path.join(output_dir, f"{company_name}_comprehensive_valuation_{timestamp}.txt")
        self.generate_comprehensive_text_report(company_info, valuation_data, market_data, peer_comparison, txt_path)
        formats['txt'] = txt_path
        
        # Image format
        img_path = os.path.join(output_dir, f"{company_name}_comprehensive_valuation_{timestamp}.png")
        self.generate_comprehensive_image_report(company_info, valuation_data, market_data, peer_comparison, img_path)
        formats['png'] = img_path
        
        return formats

    def generate_word_report(self, 
                           company_info: Dict[str, Any],
                           valuation_data: Dict[str, Any],
                           market_data: Dict[str, Any],
                           peer_comparison: List[Dict[str, Any]],
                           file_path: str) -> str:
        """Generate a detailed valuation report in DOCX format"""
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('UCaaS Company Valuation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(f'Report Date: {datetime.now().strftime("%B %d, %Y")}')
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        summary = doc.add_paragraph()
        summary.add_run('Company Overview\n').bold = True
        summary.add_run(f'Company Name: {company_info.get("name", "N/A")}\n')
        summary.add_run(f'Industry: UCaaS (Unified Communications as a Service)\n')
        summary.add_run(f'Annual Recurring Revenue (ARR): ${company_info.get("arr", 0):,.2f}\n')
        
        # Key Metrics
        doc.add_heading('Key Financial Metrics', level=1)
        metrics_table = doc.add_table(rows=1, cols=2)
        metrics_table.style = 'Table Grid'
        
        # Add header row
        header_cells = metrics_table.rows[0].cells
        header_cells[0].text = 'Metric'
        header_cells[1].text = 'Value'
        
        metrics = [
            ('Growth Rate', f'{valuation_data.get("growth_rate", 0)*100:.1f}%'),
            ('Gross Margin', f'{valuation_data.get("gross_margin", 0)*100:.1f}%'),
            ('Net Revenue Retention', f'{valuation_data.get("net_revenue_retention", 0)*100:.1f}%'),
            ('Rule of 40 Score', f'{valuation_data.get("rule_of_40", 0):.1f}'),
            ('LTV/CAC Ratio', f'{valuation_data.get("ltv_cac_ratio", 0):.2f}'),
            ('Company Valuation', f'${valuation_data.get("valuation", 0):,.2f}'),
        ]
        
        for metric, value in metrics:
            row_cells = metrics_table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
        
        # Market Analysis
        doc.add_heading('Market Analysis', level=1)
        market_para = doc.add_paragraph()
        market_para.add_run(f'Market Size: ${market_data.get("market_size", 0):,.2f}\n')
        market_para.add_run(f'Market Growth Rate: {market_data.get("market_growth", 0)*100:.1f}%\n')
        market_para.add_run(f'Competitive Position: {market_data.get("competitive_position", "N/A")}\n')
        
        # Valuation Summary
        doc.add_heading('Valuation Summary', level=1)
        valuation_para = doc.add_paragraph()
        valuation_para.add_run(f'Total Company Valuation: ${valuation_data.get("valuation", 0):,.2f}\n').bold = True
        valuation_para.add_run(f'Revenue Multiple: {valuation_data.get("revenue_multiple", 0):.2f}x\n')
        valuation_para.add_run(f'EBITDA Multiple: {valuation_data.get("ebitda_multiple", 0):.2f}x\n')
        
        doc.save(file_path)
        return file_path
    
    def generate_pdf_report(self, 
                          company_info: Dict[str, Any],
                          valuation_data: Dict[str, Any],
                          market_data: Dict[str, Any],
                          peer_comparison: List[Dict[str, Any]],
                          file_path: str) -> str:
        """Generate a PDF valuation report"""
        
        doc = SimpleDocTemplate(file_path, pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=30
        )
        
        story = []
        
        # Title
        title = Paragraph("UCaaS Company Valuation Report", title_style)
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Date
        date_style = ParagraphStyle('DateStyle', parent=styles['Normal'], alignment=TA_RIGHT)
        date_para = Paragraph(f"Report Date: {datetime.now().strftime('%B %d, %Y')}", date_style)
        story.append(date_para)
        story.append(Spacer(1, 12))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", styles['Heading2']))
        story.append(Spacer(1, 6))
        
        summary_text = f"""
        <b>Company Name:</b> {company_info.get("name", "N/A")}<br/>
        <b>Industry:</b> UCaaS (Unified Communications as a Service)<br/>
        <b>Annual Recurring Revenue (ARR):</b> ${company_info.get("arr", 0):,.2f}<br/>
        """
        story.append(Paragraph(summary_text, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Key Metrics Table
        story.append(Paragraph("Key Financial Metrics", styles['Heading2']))
        story.append(Spacer(1, 6))
        
        metrics_data = [
            ['Metric', 'Value'],
            ['Growth Rate', f'{valuation_data.get("growth_rate", 0)*100:.1f}%'],
            ['Gross Margin', f'{valuation_data.get("gross_margin", 0)*100:.1f}%'],
            ['Net Revenue Retention', f'{valuation_data.get("net_revenue_retention", 0)*100:.1f}%'],
            ['Rule of 40 Score', f'{valuation_data.get("rule_of_40", 0):.1f}'],
            ['LTV/CAC Ratio', f'{valuation_data.get("ltv_cac_ratio", 0):.2f}'],
            ['Company Valuation', f'${valuation_data.get("valuation", 0):,.2f}'],
        ]
        
        metrics_table = Table(metrics_data, colWidths=[3*inch, 2*inch])
        metrics_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 14),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(metrics_table)
        story.append(Spacer(1, 12))
        
        # Valuation Summary
        story.append(Paragraph("Valuation Summary", styles['Heading2']))
        story.append(Spacer(1, 6))
        
        valuation_text = f"""
        <b>Total Company Valuation:</b> ${valuation_data.get("valuation", 0):,.2f}<br/>
        <b>Revenue Multiple:</b> {valuation_data.get("revenue_multiple", 0):.2f}x<br/>
        <b>EBITDA Multiple:</b> {valuation_data.get("ebitda_multiple", 0):.2f}x<br/>
        """
        story.append(Paragraph(valuation_text, styles['Normal']))
        
        doc.build(story)
        return file_path
    
    def generate_text_report(self, 
                           company_info: Dict[str, Any],
                           valuation_data: Dict[str, Any],
                           market_data: Dict[str, Any],
                           peer_comparison: List[Dict[str, Any]],
                           file_path: str) -> str:
        """Generate a plain text valuation report"""
        
        report_content = f"""
UCaaS COMPANY VALUATION REPORT
{'='*50}

Report Date: {datetime.now().strftime('%B %d, %Y')}

EXECUTIVE SUMMARY
{'-'*20}
Company Name: {company_info.get("name", "N/A")}
Industry: UCaaS (Unified Communications as a Service)
Annual Recurring Revenue (ARR): ${company_info.get("arr", 0):,.2f}

KEY FINANCIAL METRICS
{'-'*25}
Growth Rate: {valuation_data.get("growth_rate", 0)*100:.1f}%
Gross Margin: {valuation_data.get("gross_margin", 0)*100:.1f}%
Net Revenue Retention: {valuation_data.get("net_revenue_retention", 0)*100:.1f}%
Rule of 40 Score: {valuation_data.get("rule_of_40", 0):.1f}
LTV/CAC Ratio: {valuation_data.get("ltv_cac_ratio", 0):.2f}

MARKET ANALYSIS
{'-'*16}
Market Size: ${market_data.get("market_size", 0):,.2f}
Market Growth Rate: {market_data.get("market_growth", 0)*100:.1f}%
Competitive Position: {market_data.get("competitive_position", "N/A")}

VALUATION SUMMARY
{'-'*18}
Total Company Valuation: ${valuation_data.get("valuation", 0):,.2f}
Revenue Multiple: {valuation_data.get("revenue_multiple", 0):.2f}x
EBITDA Multiple: {valuation_data.get("ebitda_multiple", 0):.2f}x

DISCLAIMER
{'-'*10}
This valuation report is based on the information provided and standard UCaaS industry metrics.
The valuation is an estimate and should not be considered as investment advice.
Actual market conditions and company-specific factors may affect the true valuation.

Report Generated by ValuAI - UCaaS Valuation Platform
"""
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(report_content)
        
        return file_path
    
    def generate_image_report(self, 
                            company_info: Dict[str, Any],
                            valuation_data: Dict[str, Any],
                            market_data: Dict[str, Any],
                            peer_comparison: List[Dict[str, Any]],
                            file_path: str) -> str:
        """Generate an image-based valuation report"""
        
        try:
            # Create a figure with multiple subplots
            fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(16, 12))
            fig.suptitle(f'UCaaS Valuation Report - {company_info.get("name", "Company")}', fontsize=20, fontweight='bold')
            
            # 1. Key Metrics Bar Chart
            metrics = ['Growth Rate', 'Gross Margin', 'NRR', 'Rule of 40']
            values = [
                max(0, valuation_data.get("growth_rate", 0.2) * 100),
                max(0, valuation_data.get("gross_margin", 0.7) * 100),
                max(0, valuation_data.get("net_revenue_retention", 1.1) * 100),
                max(0, valuation_data.get("rule_of_40", 40))
            ]
            
            # Ensure no NaN or infinite values
            values = [v if not (v != v or v == float('inf')) else 0 for v in values]
            
            ax1.bar(metrics, values, color=['#1f77b4', '#ff7f0e', '#2ca02c', '#d62728'])
            ax1.set_title('Key Performance Metrics', fontweight='bold')
            ax1.set_ylabel('Percentage / Score')
            plt.setp(ax1.xaxis.get_majorticklabels(), rotation=45)
            
            # Add value labels on bars
            max_value = max(values) if values else 100
            for i, v in enumerate(values):
                ax1.text(i, v + max_value*0.01, f'{v:.1f}%' if i < 3 else f'{v:.1f}', 
                        ha='center', va='bottom', fontweight='bold')
            
            # 2. Valuation Breakdown Pie Chart
            valuation = valuation_data.get("final_valuation", valuation_data.get("valuation", 5000000))
            valuation_components = {
                'Revenue Multiple': max(0, valuation * 0.4),
                'Growth Premium': max(0, valuation * 0.3),
                'Market Position': max(0, valuation * 0.2),
                'Other Factors': max(0, valuation * 0.1)
            }
            
            # Filter out zero or negative values
            valuation_components = {k: v for k, v in valuation_components.items() if v > 0}
            
            if valuation_components:
                ax2.pie(valuation_components.values(), labels=valuation_components.keys(), autopct='%1.1f%%', startangle=90)
            else:
                ax2.text(0.5, 0.5, 'Valuation Data\nNot Available', ha='center', va='center', transform=ax2.transAxes)
            ax2.set_title('Valuation Components', fontweight='bold')
            
            # 3. Financial Summary Table
            ax3.axis('tight')
            ax3.axis('off')
            
            table_data = [
                ['Metric', 'Value'],
                ['Company Valuation', f'${valuation:,.0f}'],
                ['Annual Recurring Revenue', f'${company_info.get("arr", 0):,.0f}'],
                ['Revenue Multiple', f'{valuation_data.get("revenue_multiple", valuation/(max(company_info.get("arr", 1), 1))):.2f}x'],
                ['LTV/CAC Ratio', f'{valuation_data.get("ltv_cac_ratio", 4.2):.2f}'],
                ['Market Size', f'${market_data.get("market_size", 50000000000):,.0f}']
            ]
            
            table = ax3.table(cellText=table_data, cellLoc='center', loc='center',
                             colWidths=[0.5, 0.5])
            table.auto_set_font_size(False)
            table.set_fontsize(12)
            table.scale(1.2, 1.5)
            
            # Style the header row
            for i in range(len(table_data[0])):
                table[(0, i)].set_facecolor('#4472C4')
                table[(0, i)].set_text_props(weight='bold', color='white')
            
            ax3.set_title('Financial Summary', fontweight='bold')
            
            # 4. Growth Trend (simulated)
            years = list(range(2020, 2026))
            current_arr = max(company_info.get("arr", 1000000), 1000)
            growth_rate = max(0.05, min(2.0, valuation_data.get("growth_rate", 0.3)))  # Cap growth rate
            
            projected_arr = [current_arr * (1 + growth_rate) ** (year - 2025) for year in years]
            
            ax4.plot(years, projected_arr, marker='o', linewidth=3, markersize=8, color='#2ca02c')
            ax4.fill_between(years, projected_arr, alpha=0.3, color='#2ca02c')
            ax4.set_title('ARR Growth Projection', fontweight='bold')
            ax4.set_xlabel('Year')
            ax4.set_ylabel('Annual Recurring Revenue ($)')
            ax4.grid(True, alpha=0.3)
            
            # Format y-axis to show values in millions
            ax4.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x/1e6:.1f}M'))
            
            plt.tight_layout()
            plt.savefig(file_path, dpi=300, bbox_inches='tight')
            plt.close()
            
            return file_path
            
        except Exception as e:
            print(f"Error generating image report: {e}")
            # Create a simple text-based image as fallback
            img = Image.new('RGB', (800, 600), 'white')
            draw = ImageDraw.Draw(img)
            
            try:
                font = ImageFont.truetype("arial.ttf", 24)
                title_font = ImageFont.truetype("arial.ttf", 36)
            except:
                font = ImageFont.load_default()
                title_font = ImageFont.load_default()
            
            # Draw title
            draw.text((50, 50), f'Valuation Report - {company_info.get("name", "Company")}', 
                     fill='black', font=title_font)
            
            # Draw basic info
            y_pos = 150
            valuation = valuation_data.get("final_valuation", valuation_data.get("valuation", 5000000))
            info_lines = [
                f'Company Valuation: ${valuation:,.0f}',
                f'Revenue: ${company_info.get("arr", 0):,.0f}',
                f'Selected Method: {valuation_data.get("selected_method", "DCF")}',
                f'Confidence: {valuation_data.get("confidence_score", 85):.0f}%'
            ]
            
            for line in info_lines:
                draw.text((50, y_pos), line, fill='black', font=font)
                y_pos += 50
            
            img.save(file_path)
            return file_path
        plt.savefig(file_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        return file_path
    
    def generate_comprehensive_word_report(self, 
                                         company_info: Dict[str, Any],
                                         valuation_data: Dict[str, Any],
                                         market_data: Dict[str, Any],
                                         peer_comparison: List[Dict[str, Any]],
                                         file_path: str) -> str:
        """Generate a comprehensive valuation report with all three methods in DOCX format"""
        
        doc = Document()
        
        # Add title
        title = doc.add_heading('🏆 Comprehensive UCaaS Valuation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add date and company
        date_paragraph = doc.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_paragraph.add_run(f'Report Date: {datetime.now().strftime("%B %d, %Y")}')
        
        company_paragraph = doc.add_paragraph()
        company_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_paragraph.add_run(f'Company: {company_info.get("name", "UCaaS Company")}').bold = True
        
        # Executive Summary
        doc.add_heading('📋 Executive Summary', level=1)
        
        exec_summary = doc.add_paragraph()
        exec_summary.add_run('ValuAI Recommendation: ').bold = True
        exec_summary.add_run(f'${valuation_data.get("recommended_valuation", 0):,.0f} using {valuation_data.get("recommended_method", "N/A")}\n\n')
        
        exec_summary.add_run('Confidence Level: ').bold = True
        exec_summary.add_run(f'{valuation_data.get("confidence_level", "Medium")}\n\n')
        
        exec_summary.add_run('Justification:\n').bold = True
        exec_summary.add_run(f'{valuation_data.get("justification", "Standard valuation methodology applied.")}\n\n')
        
        # Data Quality Assessment
        doc.add_heading('📊 Data Quality Assessment', level=1)
        data_quality = valuation_data.get('data_quality', {})
        
        quality_para = doc.add_paragraph()
        quality_para.add_run(f'Overall Data Quality Score: {data_quality.get("overall_score", 0)*100:.1f}%\n')
        quality_para.add_run(f'Data Completeness: {data_quality.get("data_completeness_percentage", 0):.1f}%\n')
        
        factors = data_quality.get('factors', {})
        quality_para.add_run(f'Consistency Score: {factors.get("consistency", 0)*100:.1f}%\n')
        quality_para.add_run(f'Predictability Score: {factors.get("predictability", 0)*100:.1f}%\n')
        
        # Three Valuation Methods
        doc.add_heading('🔍 Three Valuation Methods Analyzed', level=1)
        
        # Method 1: DCF Valuation
        doc.add_heading('💼 1. DCF Valuation (Discounted Cash Flow)', level=2)
        dcf_para = doc.add_paragraph()
        dcf_para.add_run('Valuation Result: ').bold = True
        dcf_para.add_run(f'${valuation_data.get("dcf_valuation", 0):,.0f}\n')
        dcf_para.add_run('Methodology: Projects future cash flows over 5-year horizon with terminal value\n')
        dcf_para.add_run('Best For: Companies with predictable revenue and stable cost structure\n')
        
        # Method 2: UCaaS-Specific Metrics
        doc.add_heading('📈 2. UCaaS-Specific Metrics', level=2)
        ucaas_para = doc.add_paragraph()
        ucaas_para.add_run('Valuation Result: ').bold = True
        ucaas_para.add_run(f'${valuation_data.get("ucaas_valuation", 0):,.0f}\n')
        ucaas_para.add_run('Key Metrics Analyzed:\n')
        ucaas_para.add_run('• MRR (Monthly Recurring Revenue)\n')
        ucaas_para.add_run('• Customer Acquisition Cost (CAC)\n')
        ucaas_para.add_run('• Lifetime Value (LTV)\n')
        ucaas_para.add_run('• Net Revenue Retention (NRR)\n')
        ucaas_para.add_run('• Rule of 40 Score\n')
        
        # Method 3: AI-Powered Valuation
        doc.add_heading('🤖 3. AI-Powered Valuation', level=2)
        ai_para = doc.add_paragraph()
        ai_para.add_run('Valuation Result: ').bold = True
        ai_para.add_run(f'${valuation_data.get("ai_valuation", 0):,.0f}\n')
        ai_para.add_run('AI Analysis: Uses machine learning trained on industry data\n')
        ai_para.add_run('Considers: Growth narrative, market position, technology differentiation\n')
        ai_para.add_run('Advantage: Pattern recognition beyond traditional financial metrics\n')
        
        # Valuation Comparison Table
        doc.add_heading('📊 Valuation Method Comparison', level=1)
        
        comparison_table = doc.add_table(rows=1, cols=3)
        comparison_table.style = 'Table Grid'
        
        # Header row
        header_cells = comparison_table.rows[0].cells
        header_cells[0].text = 'Method'
        header_cells[1].text = 'Valuation'
        header_cells[2].text = 'Best Use Case'
        
        methods_data = [
            ('DCF Valuation', f'${valuation_data.get("dcf_valuation", 0):,.0f}', 'Predictable cash flows'),
            ('UCaaS Metrics', f'${valuation_data.get("ucaas_valuation", 0):,.0f}', 'Recurring revenue strength'),
            ('AI-Powered', f'${valuation_data.get("ai_valuation", 0):,.0f}', 'Complex pattern recognition')
        ]
        
        for method, valuation, use_case in methods_data:
            row_cells = comparison_table.add_row().cells
            row_cells[0].text = method
            row_cells[1].text = valuation
            row_cells[2].text = use_case
        
        # Final Recommendation
        doc.add_heading('🎯 Final Recommendation', level=1)
        
        final_para = doc.add_paragraph()
        final_para.add_run('Recommended Valuation: ').bold = True
        final_para.add_run(f'${valuation_data.get("recommended_valuation", 0):,.0f}\n\n').bold = True
        
        final_para.add_run('Selected Method: ').bold = True
        final_para.add_run(f'{valuation_data.get("recommended_method", "N/A")}\n\n')
        
        final_para.add_run('Why This Method?\n').bold = True
        final_para.add_run(f'{valuation_data.get("justification", "Standard methodology applied.")}\n\n')
        
        # Valuation Range
        valuation_range = valuation_data.get('valuation_range', {})
        if valuation_range:
            final_para.add_run('Valuation Range Analysis:\n').bold = True
            final_para.add_run(f'• Low: ${valuation_range.get("low", 0):,.0f}\n')
            final_para.add_run(f'• High: ${valuation_range.get("high", 0):,.0f}\n')
            final_para.add_run(f'• Average: ${valuation_range.get("average", 0):,.0f}\n')
        
        # Market Context
        doc.add_heading('🌍 Market Context', level=1)
        market_para = doc.add_paragraph()
        market_para.add_run(f'UCaaS Market Size: ${market_data.get("market_size", 50000000000):,.0f}\n')
        market_para.add_run(f'Market Growth Rate: {market_data.get("market_growth", 0.15)*100:.1f}% annually\n')
        market_para.add_run(f'Competitive Position: {market_data.get("competitive_position", "Average")}\n')
        
        # Disclaimer
        doc.add_heading('⚠️ Important Disclaimer', level=1)
        disclaimer_para = doc.add_paragraph()
        disclaimer_para.add_run('This valuation report is generated by ValuAI using industry-standard methodologies and AI analysis. ')
        disclaimer_para.add_run('The valuation is an estimate based on provided data and should not be considered as investment advice. ')
        disclaimer_para.add_run('Actual market conditions, competitive dynamics, and company-specific factors may significantly affect the true valuation. ')
        disclaimer_para.add_run('Please consult with qualified financial professionals for investment decisions.')
        
        doc.save(file_path)
        return file_path
    
    def generate_comprehensive_pdf_report(self, 
                                        company_info: Dict[str, Any],
                                        valuation_data: Dict[str, Any],
                                        market_data: Dict[str, Any],
                                        peer_comparison: List[Dict[str, Any]],
                                        file_path: str) -> str:
        """Generate a comprehensive PDF valuation report with all three methods"""
        
        doc = SimpleDocTemplate(file_path, pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=20,
            alignment=TA_CENTER,
            spaceAfter=30
        )
        
        story = []
        
        # Title
        title = Paragraph("🏆 Comprehensive UCaaS Valuation Report", title_style)
        story.append(title)
        story.append(Spacer(1, 12))
        
        # Company and Date
        company_style = ParagraphStyle('CompanyStyle', parent=styles['Normal'], 
                                     alignment=TA_CENTER, fontSize=14, textColor=colors.blue)
        company_para = Paragraph(f"<b>{company_info.get('name', 'UCaaS Company')}</b>", company_style)
        story.append(company_para)
        story.append(Spacer(1, 6))
        
        date_style = ParagraphStyle('DateStyle', parent=styles['Normal'], alignment=TA_RIGHT)
        date_para = Paragraph(f"Report Date: {datetime.now().strftime('%B %d, %Y')}", date_style)
        story.append(date_para)
        story.append(Spacer(1, 20))
        
        # Executive Summary
        story.append(Paragraph("📋 Executive Summary", styles['Heading2']))
        story.append(Spacer(1, 6))
        
        exec_summary = f"""
        <b>ValuAI Recommendation:</b> ${valuation_data.get("recommended_valuation", 0):,.0f} using {valuation_data.get("recommended_method", "N/A")}<br/><br/>
        <b>Confidence Level:</b> {valuation_data.get("confidence_level", "Medium")}<br/><br/>
        <b>Justification:</b><br/>
        {valuation_data.get("justification", "Standard valuation methodology applied.")}
        """
        story.append(Paragraph(exec_summary, styles['Normal']))
        story.append(Spacer(1, 15))
        
        # Three Methods Analysis
        story.append(Paragraph("🔍 Three Valuation Methods Analyzed", styles['Heading2']))
        story.append(Spacer(1, 6))
        
        methods_data = [
            ['Method', 'Valuation', 'Key Strengths'],
            ['💼 DCF Valuation', f'${valuation_data.get("dcf_valuation", 0):,.0f}', 'Fundamental analysis, time value of money'],
            ['📈 UCaaS Metrics', f'${valuation_data.get("ucaas_valuation", 0):,.0f}', 'Industry-specific, recurring revenue focus'],
            ['🤖 AI-Powered', f'${valuation_data.get("ai_valuation", 0):,.0f}', 'Pattern recognition, qualitative factors']
        ]
        
        methods_table = Table(methods_data, colWidths=[2*inch, 1.5*inch, 2.5*inch])
        methods_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.lightblue),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP')
        ]))
        
        story.append(methods_table)
        story.append(Spacer(1, 15))
        
        # Final Recommendation
        story.append(Paragraph("🎯 Final Recommendation", styles['Heading2']))
        story.append(Spacer(1, 6))
        
        recommendation_text = f"""
        <b>Recommended Valuation: ${valuation_data.get("recommended_valuation", 0):,.0f}</b><br/>
        <b>Selected Method:</b> {valuation_data.get("recommended_method", "N/A")}<br/><br/>
        <b>Why This Method?</b><br/>
        {valuation_data.get("justification", "Standard methodology applied.")}
        """
        story.append(Paragraph(recommendation_text, styles['Normal']))
        story.append(Spacer(1, 10))
        
        # Valuation Range
        valuation_range = valuation_data.get('valuation_range', {})
        if valuation_range:
            range_text = f"""
            <b>Valuation Range Analysis:</b><br/>
            • Low: ${valuation_range.get("low", 0):,.0f}<br/>
            • High: ${valuation_range.get("high", 0):,.0f}<br/>
            • Average: ${valuation_range.get("average", 0):,.0f}
            """
            story.append(Paragraph(range_text, styles['Normal']))
            story.append(Spacer(1, 15))
        
        # Data Quality
        story.append(Paragraph("📊 Data Quality Assessment", styles['Heading2']))
        story.append(Spacer(1, 6))
        
        data_quality = valuation_data.get('data_quality', {})
        quality_text = f"""
        Overall Data Quality Score: {data_quality.get("overall_score", 0)*100:.1f}%<br/>
        Data Completeness: {data_quality.get("data_completeness_percentage", 0):.1f}%<br/>
        Consistency Score: {data_quality.get("factors", {}).get("consistency", 0)*100:.1f}%<br/>
        Predictability Score: {data_quality.get("factors", {}).get("predictability", 0)*100:.1f}%
        """
        story.append(Paragraph(quality_text, styles['Normal']))
        story.append(Spacer(1, 15))
        
        # Market Context
        story.append(Paragraph("🌍 Market Context", styles['Heading2']))
        story.append(Spacer(1, 6))
        
        market_text = f"""
        UCaaS Market Size: ${market_data.get("market_size", 50000000000):,.0f}<br/>
        Market Growth Rate: {market_data.get("market_growth", 0.15)*100:.1f}% annually<br/>
        Competitive Position: {market_data.get("competitive_position", "Average")}
        """
        story.append(Paragraph(market_text, styles['Normal']))
        story.append(Spacer(1, 15))
        
        # Disclaimer
        story.append(Paragraph("⚠️ Important Disclaimer", styles['Heading2']))
        disclaimer_text = """
        This valuation report is generated by ValuAI using industry-standard methodologies and AI analysis. 
        The valuation is an estimate based on provided data and should not be considered as investment advice. 
        Actual market conditions, competitive dynamics, and company-specific factors may significantly affect the true valuation. 
        Please consult with qualified financial professionals for investment decisions.
        """
        story.append(Paragraph(disclaimer_text, styles['Normal']))
        
        doc.build(story)
        return file_path
    
    def generate_comprehensive_text_report(self, 
                                         company_info: Dict[str, Any],
                                         valuation_data: Dict[str, Any],
                                         market_data: Dict[str, Any],
                                         peer_comparison: List[Dict[str, Any]],
                                         file_path: str) -> str:
        """Generate a comprehensive plain text valuation report"""
        
        report_content = f"""
🏆 COMPREHENSIVE UCaaS VALUATION REPORT
{'='*60}

Company: {company_info.get("name", "UCaaS Company")}
Report Date: {datetime.now().strftime('%B %d, %Y')}

📋 EXECUTIVE SUMMARY
{'-'*25}
ValuAI Recommendation: ${valuation_data.get("recommended_valuation", 0):,.0f}
Selected Method: {valuation_data.get("recommended_method", "N/A")}
Confidence Level: {valuation_data.get("confidence_level", "Medium")}

Justification:
{valuation_data.get("justification", "Standard valuation methodology applied.")}

🔍 THREE VALUATION METHODS ANALYZED
{'-'*40}

💼 1. DCF VALUATION (DISCOUNTED CASH FLOW)
Valuation Result: ${valuation_data.get("dcf_valuation", 0):,.0f}
Methodology: Projects future cash flows over 5-year horizon
Best For: Companies with predictable revenue and stable costs
Advantages: Fundamental analysis, considers time value of money
Limitations: Sensitive to growth assumptions, requires reliable projections

📈 2. UCaaS-SPECIFIC METRICS VALUATION
Valuation Result: ${valuation_data.get("ucaas_valuation", 0):,.0f}
Key Metrics: MRR, CAC, LTV, NRR, Rule of 40, Churn Rate
Best For: Established SaaS companies with recurring revenue
Advantages: Industry-specific benchmarks, recurring revenue focus
Limitations: May overestimate with aggressive assumptions

🤖 3. AI-POWERED VALUATION
Valuation Result: ${valuation_data.get("ai_valuation", 0):,.0f}
AI Analysis: Machine learning trained on industry data
Considers: Growth narrative, market position, technology differentiation
Best For: Complex scenarios with rich qualitative data
Advantages: Pattern recognition beyond traditional metrics
Limitations: Less transparent, requires comprehensive data

📊 VALUATION METHOD COMPARISON
{'-'*35}
DCF Valuation:        ${valuation_data.get("dcf_valuation", 0):,.0f}
UCaaS Metrics:        ${valuation_data.get("ucaas_valuation", 0):,.0f}
AI-Powered:           ${valuation_data.get("ai_valuation", 0):,.0f}
RECOMMENDED:          ${valuation_data.get("recommended_valuation", 0):,.0f}

🎯 FINAL RECOMMENDATION
{'-'*25}
Recommended Valuation: ${valuation_data.get("recommended_valuation", 0):,.0f}
Selected Method: {valuation_data.get("recommended_method", "N/A")}

Why This Method?
{valuation_data.get("justification", "Standard methodology applied.")}

VALUATION RANGE ANALYSIS:"""

        valuation_range = valuation_data.get('valuation_range', {})
        if valuation_range:
            report_content += f"""
Low Estimate:         ${valuation_range.get("low", 0):,.0f}
High Estimate:        ${valuation_range.get("high", 0):,.0f}
Average:              ${valuation_range.get("average", 0):,.0f}
Median:               ${valuation_range.get("median", 0):,.0f}"""

        data_quality = valuation_data.get('data_quality', {})
        report_content += f"""

📊 DATA QUALITY ASSESSMENT
{'-'*30}
Overall Quality Score: {data_quality.get("overall_score", 0)*100:.1f}%
Data Completeness: {data_quality.get("data_completeness_percentage", 0):.1f}%
Consistency Score: {data_quality.get("factors", {}).get("consistency", 0)*100:.1f}%
Predictability Score: {data_quality.get("factors", {}).get("predictability", 0)*100:.1f}%

🌍 MARKET CONTEXT
{'-'*18}
UCaaS Market Size: ${market_data.get("market_size", 50000000000):,.0f}
Market Growth Rate: {market_data.get("market_growth", 0.15)*100:.1f}% annually
Competitive Position: {market_data.get("competitive_position", "Average")}

⚠️ IMPORTANT DISCLAIMER
{'-'*23}
This valuation report is generated by ValuAI using industry-standard 
methodologies and AI analysis. The valuation is an estimate based on 
provided data and should not be considered as investment advice. 
Actual market conditions, competitive dynamics, and company-specific 
factors may significantly affect the true valuation. Please consult 
with qualified financial professionals for investment decisions.

Report Generated by ValuAI - Comprehensive UCaaS Valuation Platform
{'='*60}
"""
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(report_content)
        
        return file_path
    
    def generate_comprehensive_image_report(self, 
                                          company_info: Dict[str, Any],
                                          valuation_data: Dict[str, Any],
                                          market_data: Dict[str, Any],
                                          peer_comparison: List[Dict[str, Any]],
                                          file_path: str) -> str:
        """Generate a comprehensive image-based valuation report"""
        
        # Create a larger figure for comprehensive report
        fig = plt.figure(figsize=(20, 16))
        
        # Set up the main title
        fig.suptitle(f'🏆 Comprehensive UCaaS Valuation Report - {company_info.get("name", "Company")}', 
                    fontsize=24, fontweight='bold', y=0.95)
        
        # Create a grid layout: 3x3
        gs = fig.add_gridspec(4, 3, hspace=0.3, wspace=0.3)
        
        # 1. Three Methods Comparison (top row, spanning 2 columns)
        ax1 = fig.add_subplot(gs[0, :2])
        methods = ['DCF\nValuation', 'UCaaS\nMetrics', 'AI-Powered\nValuation']
        valuations = [
            valuation_data.get("dcf_valuation", 0),
            valuation_data.get("ucaas_valuation", 0),
            valuation_data.get("ai_valuation", 0)
        ]
        
        colors_methods = ['#1f77b4', '#ff7f0e', '#2ca02c']
        bars = ax1.bar(methods, valuations, color=colors_methods, alpha=0.8)
        ax1.set_title('💼📊🤖 Three Valuation Methods Comparison', fontsize=16, fontweight='bold')
        ax1.set_ylabel('Valuation ($)', fontsize=12)
        
        # Add value labels on bars
        for bar, val in zip(bars, valuations):
            height = bar.get_height()
            ax1.text(bar.get_x() + bar.get_width()/2., height + max(valuations)*0.01,
                    f'${val:,.0f}', ha='center', va='bottom', fontweight='bold', fontsize=11)
        
        # Format y-axis
        ax1.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x/1e6:.1f}M'))
        
        # 2. Recommended Method Highlight (top right)
        ax2 = fig.add_subplot(gs[0, 2])
        ax2.axis('off')
        
        # Create a highlighted box for recommendation
        recommended_val = valuation_data.get("recommended_valuation", 0)
        recommended_method = valuation_data.get("recommended_method", "N/A")
        confidence = valuation_data.get("confidence_level", "Medium")
        
        ax2.text(0.5, 0.8, '🎯 RECOMMENDATION', ha='center', va='center', 
                fontsize=14, fontweight='bold', transform=ax2.transAxes)
        ax2.text(0.5, 0.6, f'${recommended_val:,.0f}', ha='center', va='center', 
                fontsize=18, fontweight='bold', color='red', transform=ax2.transAxes)
        ax2.text(0.5, 0.4, f'Method: {recommended_method}', ha='center', va='center', 
                fontsize=12, transform=ax2.transAxes)
        ax2.text(0.5, 0.2, f'Confidence: {confidence}', ha='center', va='center', 
                fontsize=12, transform=ax2.transAxes)
        
        # Add border
        rect = patches.Rectangle((0.1, 0.1), 0.8, 0.8, linewidth=3, 
                               edgecolor='red', facecolor='lightpink', alpha=0.3,
                               transform=ax2.transAxes)
        ax2.add_patch(rect)
        
        # 3. Data Quality Dashboard (second row, left)
        ax3 = fig.add_subplot(gs[1, 0])
        data_quality = valuation_data.get('data_quality', {})
        factors = data_quality.get('factors', {})
        
        quality_metrics = ['Overall', 'Completeness', 'Consistency', 'Predictability']
        quality_scores = [
            data_quality.get('overall_score', 0) * 100,
            factors.get('completeness', 0) * 100,
            factors.get('consistency', 0) * 100,
            factors.get('predictability', 0) * 100
        ]
        
        bars3 = ax3.barh(quality_metrics, quality_scores, color=['#ff6b6b', '#4ecdc4', '#45b7d1', '#96ceb4'])
        ax3.set_title('📊 Data Quality Scores', fontweight='bold')
        ax3.set_xlabel('Score (%)')
        ax3.set_xlim(0, 100)
        
        # Add percentage labels
        for bar, score in zip(bars3, quality_scores):
            width = bar.get_width()
            ax3.text(width + 2, bar.get_y() + bar.get_height()/2, 
                    f'{score:.1f}%', ha='left', va='center', fontweight='bold')
        
        # 4. Valuation Range Analysis (second row, middle)
        ax4 = fig.add_subplot(gs[1, 1])
        valuation_range = valuation_data.get('valuation_range', {})
        
        if valuation_range:
            range_labels = ['Low', 'Average', 'High']
            range_values = [
                valuation_range.get('low', 0),
                valuation_range.get('average', 0),
                valuation_range.get('high', 0)
            ]
            
            ax4.plot(range_labels, range_values, marker='o', linewidth=3, 
                    markersize=10, color='purple')
            ax4.fill_between(range_labels, range_values, alpha=0.3, color='purple')
            ax4.set_title('📈 Valuation Range', fontweight='bold')
            ax4.set_ylabel('Valuation ($)')
            ax4.yaxis.set_major_formatter(plt.FuncFormatter(lambda x, p: f'${x/1e6:.1f}M'))
            
            # Highlight recommended value
            recommended_val = valuation_data.get("recommended_valuation", 0)
            ax4.axhline(y=recommended_val, color='red', linestyle='--', linewidth=2, alpha=0.8)
            ax4.text(1, recommended_val, f'Recommended: ${recommended_val/1e6:.1f}M', 
                    ha='center', va='bottom', fontweight='bold', color='red')
        
        # 5. Market Context (second row, right)
        ax5 = fig.add_subplot(gs[1, 2])
        ax5.axis('off')
        
        ax5.text(0.5, 0.9, '🌍 MARKET CONTEXT', ha='center', va='center', 
                fontsize=14, fontweight='bold', transform=ax5.transAxes)
        ax5.text(0.1, 0.7, f'Market Size:', ha='left', va='center', 
                fontsize=11, fontweight='bold', transform=ax5.transAxes)
        ax5.text(0.1, 0.6, f'${market_data.get("market_size", 50000000000)/1e9:.0f}B', 
                ha='left', va='center', fontsize=11, transform=ax5.transAxes)
        ax5.text(0.1, 0.4, f'Growth Rate:', ha='left', va='center', 
                fontsize=11, fontweight='bold', transform=ax5.transAxes)
        ax5.text(0.1, 0.3, f'{market_data.get("market_growth", 0.15)*100:.1f}%', 
                ha='left', va='center', fontsize=11, transform=ax5.transAxes)
        ax5.text(0.1, 0.1, f'Position: {market_data.get("competitive_position", "Average")}', 
                ha='left', va='center', fontsize=11, transform=ax5.transAxes)
        
        # 6. Method Strengths Comparison (third row, spanning all columns)
        ax6 = fig.add_subplot(gs[2, :])
        ax6.axis('off')
        
        # Create a comparison table
        method_strengths = [
            ['DCF Valuation', '• Fundamental cash flow analysis\n• Time value of money\n• Widely accepted'],
            ['UCaaS Metrics', '• Industry-specific benchmarks\n• Recurring revenue focus\n• Customer retention analysis'],
            ['AI-Powered', '• Pattern recognition\n• Qualitative factor integration\n• Market sentiment analysis']
        ]
        
        ax6.text(0.5, 0.9, '🔍 Method Strengths & Applications', ha='center', va='center', 
                fontsize=16, fontweight='bold', transform=ax6.transAxes)
        
        for i, (method, strengths) in enumerate(method_strengths):
            x_pos = 0.15 + i * 0.3
            ax6.text(x_pos, 0.7, method, ha='center', va='center', 
                    fontsize=12, fontweight='bold', transform=ax6.transAxes)
            ax6.text(x_pos, 0.4, strengths, ha='center', va='center', 
                    fontsize=10, transform=ax6.transAxes)
            
            # Add background color
            rect = patches.Rectangle((x_pos-0.12, 0.15), 0.24, 0.65, 
                                   linewidth=1, edgecolor='gray', 
                                   facecolor=colors_methods[i], alpha=0.1,
                                   transform=ax6.transAxes)
            ax6.add_patch(rect)
        
        # 7. Justification Text (bottom row)
        ax7 = fig.add_subplot(gs[3, :])
        ax7.axis('off')
        
        justification = valuation_data.get("justification", "Standard valuation methodology applied.")
        ax7.text(0.5, 0.8, '💡 WHY THIS VALUATION METHOD?', ha='center', va='center', 
                fontsize=16, fontweight='bold', transform=ax7.transAxes)
        
        # Wrap text for better display
        import textwrap
        wrapped_text = textwrap.fill(justification, width=120)
        ax7.text(0.5, 0.4, wrapped_text, ha='center', va='center', 
                fontsize=12, transform=ax7.transAxes, 
                bbox=dict(boxstyle="round,pad=0.5", facecolor="lightblue", alpha=0.8))
        
        # Add footer
        ax7.text(0.5, 0.05, 'Generated by ValuAI - Comprehensive UCaaS Valuation Platform', 
                ha='center', va='center', fontsize=10, style='italic', 
                transform=ax7.transAxes)
        
        plt.savefig(file_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.close()
        
        return file_path
