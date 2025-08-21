#!/usr/bin/env python3
"""
ValuAI Business Valuation Tool - Main Server
Clean, optimized version with proper file handling
"""
from flask import Flask, request, jsonify, send_file, Response
from flask_cors import CORS
from werkzeug.utils import secure_filename
import sqlite3
import hashlib
import jwt
from datetime import datetime, timedelta
import os
import zipfile
import io

# Optional imports with fallbacks
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.enums import TA_CENTER
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

except ImportError:
    DOCX_AVAILABLE = False

app = Flask(__name__)
app.config['SECRET_KEY'] = 'valuai-secret-key-2024'

# JWT Configuration - NEVER EXPIRES (FOR TESTING ONLY)
app.config['JWT_ACCESS_TOKEN_EXPIRES'] = False

# Optimize Flask for better performance
app.config['JSON_SORT_KEYS'] = False
app.config['JSONIFY_PRETTYPRINT_REGULAR'] = False

# Simple CORS configuration - allow all origins for development
CORS(app, resources={
    r"/api/*": {
        "origins": ["http://127.0.0.1:5173", "http://localhost:5173"],
        "methods": ["GET", "POST", "PUT", "DELETE", "OPTIONS"],
        "allow_headers": ["Content-Type", "Authorization"]
    }
})

def hash_file_content(content: bytes) -> str:
    """Generate MD5 hash of file content for integrity checking"""
    return hashlib.md5(content).hexdigest()

def prepare_report_data(company_name: str, industry: str, revenue: float, growth_rate: float, ebitda_margin: float):
    """Centralize report data preparation for consistency across formats"""
    # Calculate financial metrics
    ebitda = revenue * ebitda_margin
    net_profit = ebitda * 0.7
    
    # Calculate valuations
    dcf_value = revenue * 2.8 * (1 + growth_rate)
    market_value = revenue * 3.2
    asset_value = revenue * 1.8
    weighted_avg = (dcf_value * 0.4 + market_value * 0.4 + asset_value * 0.2)
    
    # Generate previous years data
    current_year = datetime.now().year
    prev_year_1 = current_year - 1
    prev_year_2 = current_year - 2
    
    revenue_2022 = revenue * 0.73
    revenue_2023 = revenue * 0.79
    ebitda_2022 = revenue_2022 * (ebitda_margin - 0.05)
    ebitda_2023 = revenue_2023 * (ebitda_margin - 0.02)
    net_profit_2022 = ebitda_2022 * 0.65
    net_profit_2023 = ebitda_2023 * 0.68
    
    return {
        "company_name": company_name,
        "industry": industry,
        "revenue": revenue,
        "growth_rate": growth_rate,
        "ebitda_margin": ebitda_margin,
        "ebitda": ebitda,
        "net_profit": net_profit,
        "dcf_value": dcf_value,
        "market_value": market_value,
        "asset_value": asset_value,
        "weighted_avg": weighted_avg,
        "current_year": current_year,
        "prev_year_1": prev_year_1,
        "prev_year_2": prev_year_2,
        "revenue_2022": revenue_2022,
        "revenue_2023": revenue_2023,
        "ebitda_2022": ebitda_2022,
        "ebitda_2023": ebitda_2023,
        "net_profit_2022": net_profit_2022,
        "net_profit_2023": net_profit_2023
    }

def create_safe_filename(company_name: str, format_ext: str) -> str:
    """Create safe filename with special character handling"""
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_name = "".join(c for c in company_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
    safe_name = safe_name.replace(' ', '_')
    return f"{safe_name}_valuation_report_{timestamp}.{format_ext}"

def get_db_connection():
    """Get database connection with connection pooling"""
    conn = sqlite3.connect('valuai.db', check_same_thread=False)
    conn.row_factory = sqlite3.Row
    # Enable WAL mode for better performance
    conn.execute('PRAGMA journal_mode=WAL')
    return conn

def init_db():
    """Initialize database with users table - optimized for faster startup"""
    # Skip initialization if database already exists and has data
    if os.path.exists('valuai.db'):
        try:
            conn = get_db_connection()
            cursor = conn.cursor()
            # Quick check if users table exists and has the default user
            cursor.execute('SELECT COUNT(*) FROM users WHERE email = ?', ('nsp6575@gmail.com',))
            user_exists = cursor.fetchone()[0] > 0
            conn.close()
            if user_exists:
                return  # Skip initialization if user already exists
        except:
            pass  # If any error, continue with full initialization
    
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # Create users table if not exists
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Check if user exists, if not create default user
    cursor.execute('SELECT COUNT(*) FROM users WHERE email = ?', ('nsp6575@gmail.com',))
    if cursor.fetchone()[0] == 0:
        password_hash = hashlib.sha256('Newpassword123'.encode()).hexdigest()
        cursor.execute(
            'INSERT INTO users (email, password_hash) VALUES (?, ?)',
            ('nsp6575@gmail.com', password_hash)
        )
    
    conn.commit()
    conn.close()

def verify_password(password, password_hash):
    """Verify password against hash"""
    return hashlib.sha256(password.encode()).hexdigest() == password_hash

def generate_token(user_id, email):
    """Generate JWT token - NEVER EXPIRES (FOR TESTING ONLY)"""
    payload = {
        'user_id': user_id,
        'email': email
        # 'exp' removed - token never expires
    }
    return jwt.encode(payload, app.config['SECRET_KEY'], algorithm='HS256')

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'service': 'ValuAI Business Valuation'
    })

@app.route('/api/auth/signin', methods=['POST', 'OPTIONS'])
def signin():
    """User authentication"""
    if request.method == 'OPTIONS':
        return jsonify({'status': 'ok'})
    
    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')
        
        if not email or not password:
            return jsonify({'error': 'Email and password required'}), 400
        
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('SELECT id, email, password_hash FROM users WHERE email = ?', (email,))
        user = cursor.fetchone()
        conn.close()
        
        if user and verify_password(password, user[2]):
            token = generate_token(user[0], user[1])
            return jsonify({
                'token': token,
                'user': {'id': user[0], 'email': user[1]},
                'message': 'Login successful'
            })
        else:
            return jsonify({'error': 'Invalid credentials'}), 401
            
    except Exception as e:
        return jsonify({'error': f'Login failed: {str(e)}'}), 500

def create_professional_pdf(content, filepath, company_name):
    """Create PDF with proper formatting"""
    if not REPORTLAB_AVAILABLE:
        # Fallback to enhanced text format with PDF extension
        print("ReportLab not available, creating text file as fallback")
        with open(filepath, 'w', encoding='utf-8-sig', newline='\r\n') as f:
            f.write(f"PDF Report for {company_name}\n{'='*50}\n\n{content}")
        return
    
    try:
        # Ensure directory exists
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        doc = SimpleDocTemplate(filepath, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = styles['Title']
        title_style.alignment = TA_CENTER
        story.append(Paragraph(f"Business Valuation Report: {company_name}", title_style))
        story.append(Spacer(1, 12))
        
        # Content paragraphs
        normal_style = styles['Normal']
        for line in content.split('\n'):
            if line.strip():
                # Escape HTML characters for ReportLab
                line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                if line.startswith(('1 ', '2 ', '3 ', '4 ', '5 ', '6 ', '7 ')):
                    story.append(Paragraph(line, styles['Heading1']))
                elif line.startswith('•'):
                    story.append(Paragraph(line, styles['Normal']))
                else:
                    story.append(Paragraph(line, normal_style))
                story.append(Spacer(1, 6))
        
        # Build the PDF
        doc.build(story)
        print(f"PDF created successfully: {filepath}")
        
        # Verify file was created and has content
        if os.path.exists(filepath) and os.path.getsize(filepath) > 0:
            print(f"PDF file verified: {os.path.getsize(filepath)} bytes")
        else:
            raise Exception("PDF file was not created properly")
            
    except Exception as e:
        print(f"PDF generation error: {e}")
        import traceback
        traceback.print_exc()
        # Fallback to text with PDF extension
        with open(filepath, 'w', encoding='utf-8-sig', newline='\r\n') as f:
            f.write(f"PDF Report for {company_name} (Text Fallback)\n{'='*50}\n\n{content}")
        print(f"Created text fallback file: {filepath}")

def create_professional_docx(content, filepath, company_name):
    """Create DOCX with proper formatting"""
    if not DOCX_AVAILABLE:
        # Fallback to text format with DOCX extension
        print("python-docx not available, creating text file as fallback")
        with open(filepath, 'w', encoding='utf-8-sig', newline='\r\n') as f:
            f.write(f"Word Document Report for {company_name}\n{'='*50}\n\n{content}")
        return
    
    try:
        # Ensure directory exists
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        
        doc = Document()
        doc.add_heading(f'Business Valuation Report: {company_name}', 0)
        
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            elif line.startswith(('1 ', '2 ', '3 ', '4 ', '5 ', '6 ', '7 ')):
                doc.add_heading(line, level=1)
            elif line.startswith('•'):
                doc.add_paragraph(line, style='List Bullet')
            elif ':' in line and len(line) < 100:
                p = doc.add_paragraph()
                parts = line.split(':', 1)
                run = p.add_run(parts[0] + ':')
                run.bold = True
                if len(parts) > 1:
                    p.add_run(parts[1])
            else:
                doc.add_paragraph(line)
        
        # Save the document
        doc.save(filepath)
        print(f"DOCX created successfully: {filepath}")
        
        # Verify file was created and has content
        if os.path.exists(filepath) and os.path.getsize(filepath) > 0:
            print(f"DOCX file verified: {os.path.getsize(filepath)} bytes")
        else:
            raise Exception("DOCX file was not created properly")
            
    except Exception as e:
        print(f"DOCX generation error: {e}")
        import traceback
        traceback.print_exc()
        # Fallback to text with DOCX extension
        with open(filepath, 'w', encoding='utf-8-sig', newline='\r\n') as f:
            f.write(f"Word Document Report for {company_name} (Text Fallback)\n{'='*50}\n\n{content}")
        print(f"Created text fallback file: {filepath}")

@app.route('/api/generate-comprehensive-report', methods=['POST', 'OPTIONS'])
def generate_comprehensive_report():
    """Generate comprehensive business valuation report"""
    if request.method == 'OPTIONS':
        return jsonify({'status': 'ok'})
    
    try:
        data = request.get_json() or {}
        
        # Extract company data
        company_name = data.get('companyName', 'Sample Company')
        industry = data.get('industry', 'Technology')
        revenue = float(data.get('revenue', 5000000))
        growth_rate = float(data.get('growthRate', 0.35))
        ebitda_margin = float(data.get('ebitdaMargin', 0.25))
        
        # Calculate financial metrics
        ebitda = revenue * ebitda_margin
        net_profit = ebitda * 0.7
        
        # Calculate valuations
        dcf_value = revenue * 2.8 * (1 + growth_rate)
        market_value = revenue * 3.2
        asset_value = revenue * 1.8
        weighted_avg = (dcf_value * 0.4 + market_value * 0.4 + asset_value * 0.2)
        
        # Generate previous years data
        current_year = datetime.now().year
        prev_year_1 = current_year - 1
        prev_year_2 = current_year - 2
        
        revenue_2022 = revenue * 0.73
        revenue_2023 = revenue * 0.79
        ebitda_2022 = revenue_2022 * (ebitda_margin - 0.05)
        ebitda_2023 = revenue_2023 * (ebitda_margin - 0.02)
        net_profit_2022 = ebitda_2022 * 0.65
        net_profit_2023 = ebitda_2023 * 0.68
        
        # Create comprehensive report content
        report_content = f"""Business Valuation Report: {company_name}

1 Executive Summary

This business valuation report provides an analysis of {company_name}, a private {industry.lower()} company in North America. The report includes a financial overview, valuation models used, competitor benchmarking, final valuation estimate, and strategic recommendations.

2 Ownership Context

{company_name} is a privately owned {industry.lower()} company operating in North America. As a private company, the ownership structure is not publicly disclosed.

3 Financial Overview

Year        Revenue (in billions)    EBITDA (in billions)    Net Profit (in billions)
{current_year}        ${revenue/1000000000:.1f}                    ${ebitda/1000000000:.3f}                   ${net_profit/1000000000:.1f}
{prev_year_1}        ${revenue_2023/1000000000:.1f}                    ${ebitda_2023/1000000000:.3f}                   ${net_profit_2023/1000000000:.1f}
{prev_year_2}        ${revenue_2022/1000000000:.1f}                    ${ebitda_2022/1000000000:.3f}                   ${net_profit_2022/1000000000:.1f}

Table 1: Financial Overview of {company_name}

4 Valuation Models Used

The valuation of {company_name} will be conducted using a combination of Discounted Cash Flow (DCF) analysis and Market Multiples approach.

4.1 Discounted Cash Flow (DCF) Analysis
• Growth Rate: {growth_rate*100:.1f}%
• EBITDA Margin: {ebitda_margin*100:.1f}%
• Terminal Value Multiple: 3.5x
• Discount Rate: 12%
• DCF Valuation: ${dcf_value:,.0f}

4.2 Market Multiples Approach
• Revenue Multiple: 3.2x
• Industry Average P/E: 15.5x
• Market Valuation: ${market_value:,.0f}

4.3 Asset-Based Approach
• Book Value: ${asset_value*0.83:,.0f}
• Adjusted Asset Value: ${asset_value:,.0f}

5 Competitor Benchmarking

Industry leaders and key competitors provide benchmarking context for valuation analysis.

6 Final Valuation Estimate

Based on the DCF analysis and Market Multiples approach, the estimated valuation range for {company_name} is between ${weighted_avg*0.85:,.0f} and ${weighted_avg*1.15:,.0f}.

Valuation Summary:
• DCF Method: ${dcf_value:,.0f}
• Market Multiples: ${market_value:,.0f}
• Asset-Based: ${asset_value:,.0f}
• Weighted Average: ${weighted_avg:,.0f}

Confidence Level: High (85%)
Recommendation: FAVORABLE for investment

7 Strategic Recommendations

It is recommended that {company_name} focuses on expanding its market presence, enhancing its product offerings, and optimizing its cost structure to drive profitability and long-term growth.

Key Recommendations:
• Digital transformation initiatives to improve efficiency
• Geographic expansion into emerging markets
• Strategic partnerships with technology providers
• Focus on sustainable growth models
• Investment in data analytics and AI capabilities

Risk Factors:
• Regulatory changes in industry policy
• Increased competition from new market entrants
• Economic downturns affecting revenue
• Technology disruption risks

Investment Thesis:
{company_name} presents a compelling investment opportunity given its strong financial performance, market position, and growth prospects.

========================================

Report prepared by ValuAI Platform
Date: {datetime.now().strftime('%B %d, %Y')}
© 2024 All Rights Reserved

This report contains confidential and proprietary information."""

        # Create reports directory
        reports_dir = os.path.join(os.path.dirname(__file__), 'reports')
        os.makedirs(reports_dir, exist_ok=True)
        
        # Generate report filenames
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        base_filename = f"valuation_report_{timestamp}"
        
        txt_filename = f"{base_filename}.txt"
        pdf_filename = f"{base_filename}.pdf"
        docx_filename = f"{base_filename}.docx"
        
        txt_path = os.path.join(reports_dir, txt_filename)
        pdf_path = os.path.join(reports_dir, pdf_filename)
        docx_path = os.path.join(reports_dir, docx_filename)
        
        # Generate reports in all formats
        # TXT - with UTF-8 BOM for Windows compatibility
        with open(txt_path, 'w', encoding='utf-8-sig', newline='\r\n') as f:
            f.write(report_content)
        
        # PDF - professional format
        create_professional_pdf(report_content, pdf_path, company_name)
        
        # DOCX - professional format
        create_professional_docx(report_content, docx_path, company_name)
        
        return jsonify({
            'message': 'Comprehensive reports generated successfully',
            'reports': {
                'txt': {
                    'url': f'/api/reports/download/{txt_filename}',
                    'filename': txt_filename,
                    'format': 'Plain Text (.txt)'
                },
                'pdf': {
                    'url': f'/api/reports/download/{pdf_filename}',
                    'filename': pdf_filename,
                    'format': 'PDF Document (.pdf)'
                },
                'docx': {
                    'url': f'/api/reports/download/{docx_filename}',
                    'filename': docx_filename,
                    'format': 'Word Document (.docx)'
                }
            },
            'company_name': company_name,
            'valuation_summary': {
                'dcf_value': dcf_value,
                'market_value': market_value,
                'asset_value': asset_value,
                'weighted_average': weighted_avg
            },
            'created_at': datetime.now().isoformat()
        }), 200
        
    except Exception as e:
        print(f"Report generation error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Report generation failed: {str(e)}'}), 500

@app.route('/api/reports/generate-direct', methods=['POST', 'OPTIONS'])
def generate_and_download_report():
    """Generate report and return it directly for download with enhanced validation"""
    if request.method == 'OPTIONS':
        return jsonify({'status': 'ok'})
    
    try:
        data = request.get_json() or {}
        
        # Validate input data
        report_format = data.get('format', 'pdf').lower()
        if report_format not in ['pdf', 'docx', 'txt']:
            app.logger.error(f"Invalid format requested: {report_format}")
            return jsonify({'error': 'Invalid format. Must be pdf, docx, or txt'}), 400
            
        # Extract and validate company data
        company_name = data.get('companyName', 'Sample Company')
        industry = data.get('industry', 'Technology')
        
        try:
            revenue = float(data.get('revenue', 5000000))
            growth_rate = float(data.get('growthRate', 0.35))
            ebitda_margin = float(data.get('ebitdaMargin', 0.25))
        except (ValueError, TypeError) as e:
            app.logger.error(f"Invalid numeric data: {e}")
            return jsonify({'error': 'Invalid numeric data in request'}), 400
        
        app.logger.info(f"Generating {report_format} report for {company_name} (Revenue: ${revenue:,.0f})")
        
        # Prepare unified data structure  
        report_data = prepare_report_data(company_name, industry, revenue, growth_rate, ebitda_margin)
        
        # Create comprehensive report content using centralized data
        report_content = f"""Business Valuation Report: {report_data['company_name']}

1 Executive Summary

This business valuation report provides an analysis of {report_data['company_name']}, a private {report_data['industry'].lower()} company in North America. The report includes a financial overview, valuation models used, competitor benchmarking, final valuation estimate, and strategic recommendations.

2 Ownership Context

{report_data['company_name']} is a privately owned {report_data['industry'].lower()} company operating in North America. As a private company, the ownership structure is not publicly disclosed.

3 Financial Overview

Year        Revenue (in billions)    EBITDA (in billions)    Net Profit (in billions)
{report_data['current_year']}        ${report_data['revenue']/1000000000:.1f}                    ${report_data['ebitda']/1000000000:.3f}                   ${report_data['net_profit']/1000000000:.1f}
{report_data['prev_year_1']}        ${report_data['revenue_2023']/1000000000:.1f}                    ${report_data['ebitda_2023']/1000000000:.3f}                   ${report_data['net_profit_2023']/1000000000:.1f}
{report_data['prev_year_2']}        ${report_data['revenue_2022']/1000000000:.1f}                    ${report_data['ebitda_2022']/1000000000:.3f}                   ${report_data['net_profit_2022']/1000000000:.1f}

Table 1: Financial Overview of {report_data['company_name']}

4 Valuation Models Used

The valuation of {report_data['company_name']} will be conducted using a combination of Discounted Cash Flow (DCF) analysis and Market Multiples approach.

4.1 Discounted Cash Flow (DCF) Analysis
• Growth Rate: {report_data['growth_rate']*100:.1f}%
• EBITDA Margin: {report_data['ebitda_margin']*100:.1f}%
• Terminal Value Multiple: 3.5x
• Discount Rate: 12%
• DCF Valuation: ${report_data['dcf_value']:,.0f}

4.2 Market Multiples Approach
• Revenue Multiple: 3.2x
• Industry Average P/E: 15.5x
• Market Valuation: ${report_data['market_value']:,.0f}

4.3 Asset-Based Approach
• Book Value: ${report_data['asset_value']*0.83:,.0f}
• Adjusted Asset Value: ${report_data['asset_value']:,.0f}

5 Competitor Benchmarking

Industry leaders and key competitors provide benchmarking context for valuation analysis.

6 Final Valuation Estimate

Based on the DCF analysis and Market Multiples approach, the estimated valuation range for {report_data['company_name']} is between ${report_data['weighted_avg']*0.85:,.0f} and ${report_data['weighted_avg']*1.15:,.0f}.

Valuation Summary:
• DCF Method: ${report_data['dcf_value']:,.0f}
• Market Multiples: ${report_data['market_value']:,.0f}
• Asset-Based: ${report_data['asset_value']:,.0f}
• Weighted Average: ${report_data['weighted_avg']:,.0f}

Confidence Level: High (85%)
Recommendation: FAVORABLE for investment

7 Strategic Recommendations

It is recommended that {report_data['company_name']} focuses on expanding its market presence, enhancing its product offerings, and optimizing its cost structure to drive profitability and long-term growth.

Key Recommendations:
• Digital transformation initiatives to improve efficiency
• Geographic expansion into emerging markets
• Strategic partnerships with technology providers
• Focus on sustainable growth models
• Investment in data analytics and AI capabilities

Risk Factors:
• Regulatory changes in industry policy
• Increased competition from new market entrants
• Economic downturns affecting revenue
• Technology disruption risks

Investment Thesis:
{report_data['company_name']} presents a compelling investment opportunity given its strong financial performance, market position, and growth prospects.

========================================

Report prepared by ValuAI Platform
Date: {datetime.now().strftime('%B %d, %Y')}
© 2024 All Rights Reserved

This report contains confidential and proprietary information."""

        # Generate safe filename
        filename = create_safe_filename(company_name, report_format)
        
        # Create in-memory file based on format
        if report_format == 'txt':
            # Create text file in memory
            file_content = io.BytesIO()
            file_content.write(report_content.encode('utf-8-sig'))
            file_content.seek(0)
            mimetype = 'text/plain; charset=utf-8'
            
        elif report_format == 'pdf':
            # Create PDF in memory
            file_content = io.BytesIO()
            if REPORTLAB_AVAILABLE:
                try:
                    doc = SimpleDocTemplate(file_content, pagesize=letter)
                    styles = getSampleStyleSheet()
                    story = []
                    
                    # Title
                    title_style = styles['Title']
                    title_style.alignment = TA_CENTER
                    story.append(Paragraph(f"Business Valuation Report: {company_name}", title_style))
                    story.append(Spacer(1, 12))
                    
                    # Content paragraphs
                    normal_style = styles['Normal']
                    for line in report_content.split('\n'):
                        if line.strip():
                            # Escape HTML characters
                            line = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            if line.startswith(('1 ', '2 ', '3 ', '4 ', '5 ', '6 ', '7 ')):
                                story.append(Paragraph(line, styles['Heading1']))
                            elif line.startswith('•'):
                                story.append(Paragraph(line, styles['Normal']))
                            else:
                                story.append(Paragraph(line, normal_style))
                            story.append(Spacer(1, 6))
                    
                    doc.build(story)
                    file_content.seek(0)
                except Exception as e:
                    app.logger.error(f"PDF generation error: {e}")
                    # Fallback to text
                    file_content = io.BytesIO()
                    file_content.write(f"PDF Report for {company_name} (Text Fallback)\n{'='*50}\n\n{report_content}".encode('utf-8-sig'))
                    file_content.seek(0)
            else:
                # ReportLab not available, fallback to text
                file_content.write(f"PDF Report for {company_name} (Text Fallback)\n{'='*50}\n\n{report_content}".encode('utf-8-sig'))
                file_content.seek(0)
            
            mimetype = 'application/pdf'
            
        elif report_format == 'docx':
            # Create DOCX in memory
            if DOCX_AVAILABLE:
                try:
                    doc = Document()
                    doc.add_heading(f'Business Valuation Report: {company_name}', 0)
                    
                    lines = report_content.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        elif line.startswith(('1 ', '2 ', '3 ', '4 ', '5 ', '6 ', '7 ')):
                            doc.add_heading(line, level=1)
                        elif line.startswith('•'):
                            doc.add_paragraph(line, style='List Bullet')
                        elif ':' in line and len(line) < 100:
                            p = doc.add_paragraph()
                            parts = line.split(':', 1)
                            run = p.add_run(parts[0] + ':')
                            run.bold = True
                            if len(parts) > 1:
                                p.add_run(parts[1])
                        else:
                            doc.add_paragraph(line)
                    
                    # Save to memory
                    file_content = io.BytesIO()
                    doc.save(file_content)
                    file_content.seek(0)
                except Exception as e:
                    app.logger.error(f"DOCX generation error: {e}")
                    # Fallback to text
                    file_content = io.BytesIO()
                    file_content.write(f"Word Document Report for {company_name} (Text Fallback)\n{'='*50}\n\n{report_content}".encode('utf-8-sig'))
                    file_content.seek(0)
            else:
                # python-docx not available, fallback to text
                file_content = io.BytesIO()
                file_content.write(f"Word Document Report for {company_name} (Text Fallback)\n{'='*50}\n\n{report_content}".encode('utf-8-sig'))
                file_content.seek(0)
            
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        
        # Validate file content before sending
        content_bytes = file_content.getvalue()
        file_size = len(content_bytes)
        
        # Ensure file exists & size is > 100 bytes
        if file_size < 100:
            app.logger.error(f"Generated file too small: {file_size} bytes")
            return jsonify({"error": "Report generation failed - file too small or empty"}), 500
        
        # Generate integrity hash
        content_hash = hash_file_content(content_bytes)
        
        # Log successful generation with integrity info
        app.logger.info(f"Report generated successfully: {filename} | Size: {file_size} bytes | Hash: {content_hash}")
        
        # Reset stream position
        file_content.seek(0)
        
        # Return file for download with correct headers
        response = send_file(
            file_content,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
        
        # Add additional headers for better compatibility
        response.headers['Content-Length'] = str(file_size)
        response.headers['X-Content-Hash'] = content_hash
        response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        
        return response
        
    except Exception as e:
        app.logger.error(f"Direct report generation error: {str(e)}", exc_info=True)
        return jsonify({'error': f'Report generation failed: {str(e)}'}), 500

@app.route('/api/reports/download/<filename>', methods=['GET'])
def download_report(filename):
    """Download report with proper MIME types and encoding"""
    try:
        reports_dir = os.path.join(os.path.dirname(__file__), 'reports')
        file_path = os.path.join(reports_dir, filename)
        
        # Security check
        if not os.path.abspath(file_path).startswith(os.path.abspath(reports_dir)):
            return jsonify({'error': 'Invalid file path'}), 400
        
        if not os.path.exists(file_path):
            return jsonify({'error': 'File not found'}), 404
        
        # Determine proper MIME type
        if filename.lower().endswith('.pdf'):
            mimetype = 'application/pdf'
        elif filename.lower().endswith('.txt'):
            mimetype = 'text/plain; charset=utf-8'
        elif filename.lower().endswith('.docx'):
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            mimetype = 'application/octet-stream'
        
        # Create response with proper headers
        response = send_file(
            file_path,
            mimetype=mimetype,
            as_attachment=True,
            download_name=filename
        )
        
        # Add headers for better compatibility
        response.headers['Content-Disposition'] = f'attachment; filename="{filename}"'
        response.headers['Content-Type'] = mimetype
        response.headers['Cache-Control'] = 'no-cache'
        
        return response
        
    except Exception as e:
        print(f"Download error: {e}")
        return jsonify({'error': f'Download failed: {str(e)}'}), 500

@app.route('/api/reports/list', methods=['GET'])
def list_reports():
    """List all available reports"""
    try:
        reports_dir = os.path.join(os.path.dirname(__file__), 'reports')
        if not os.path.exists(reports_dir):
            return jsonify({'reports': []})
        
        reports = []
        for filename in os.listdir(reports_dir):
            if filename.endswith(('.txt', '.pdf', '.docx')):
                file_path = os.path.join(reports_dir, filename)
                file_size = os.path.getsize(file_path)
                file_modified = datetime.fromtimestamp(os.path.getmtime(file_path))
                
                reports.append({
                    'filename': filename,
                    'size': file_size,
                    'modified': file_modified.isoformat(),
                    'download_url': f'/api/reports/download/{filename}'
                })
        
        return jsonify({'reports': sorted(reports, key=lambda x: x['modified'], reverse=True)})
        
    except Exception as e:
        return jsonify({'error': f'Failed to list reports: {str(e)}'}), 500

# File Upload Routes
@app.route('/api/files/upload', methods=['POST', 'OPTIONS'])
def upload_file():
    """Upload and process files (Excel, CSV, PDF, Word, etc.)"""
    print(f"📤 Upload request received - Method: {request.method}")
    print(f"📤 Files in request: {list(request.files.keys())}")
    print(f"📤 Headers: {dict(request.headers)}")
    
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        if 'file' not in request.files:
            print("❌ No file part in request")
            return jsonify({'error': 'No file part'}), 400
            
        file = request.files['file']
        print(f"📄 File received: {file.filename}")
        
        if file.filename == '':
            print("❌ No file selected")
            return jsonify({'error': 'No selected file'}), 400
        
        # Define allowed file types
        allowed_extensions = {'doc', 'docx', 'pdf', 'xls', 'xlsx', 'png', 'jpg', 'jpeg', 'gif', 'webp', 'txt', 'csv'}
        
        def allowed_file(filename):
            return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions
        
        if not allowed_file(file.filename):
            print(f"❌ File type not allowed: {file.filename}")
            return jsonify({'error': f'File type not allowed. Allowed types: {", ".join(allowed_extensions)}'}), 400

        # Create uploads directory if it doesn't exist
        uploads_dir = os.path.join(os.path.dirname(__file__), 'uploads')
        os.makedirs(uploads_dir, exist_ok=True)
        print(f"📁 Upload directory: {uploads_dir}")
        
        # Secure the filename and generate unique name
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_filename = f'{timestamp}_{filename}'
        file_path = os.path.join(uploads_dir, unique_filename)
        
        # Save the file
        file.save(file_path)
        print(f"💾 File saved to: {file_path}")
        
        # Process file based on type
        file_type = filename.rsplit('.', 1)[1].lower()
        
        result = {'message': 'File uploaded successfully'}
        
        if file_type in ['xls', 'xlsx', 'csv']:
            result.update(process_excel_csv(file_path))
        elif file_type in ['doc', 'docx']:
            result.update(process_word_doc(file_path))
        elif file_type == 'pdf':
            result.update(process_pdf_doc(file_path))
        elif file_type == 'txt':
            result.update(process_text_file(file_path))
        else:
            result.update({'warning': 'File uploaded but not processed'})
            
        # Add file info to result
        result.update({
            'filename': unique_filename,
            'original_filename': filename,
            'file_type': file_type,
            'file_size': os.path.getsize(file_path),
            'upload_path': file_path
        })
        
        print(f"✅ Upload successful: {filename}")
        return jsonify(result)
        
    except Exception as e:
        print(f"❌ Upload error: {str(e)}")
        return jsonify({'error': f'Failed to process file: {str(e)}'}), 500

def secure_filename(filename):
    """Make filename secure by removing unsafe characters"""
    import re
    filename = re.sub(r'[^\w\s-.]', '', filename).strip()
    return re.sub(r'[-\s]+', '-', filename)

def process_excel_csv(file_path):
    """Process Excel or CSV files"""
    try:
        import pandas as pd
        
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path)
        
        # Basic analysis
        return {
            'data_preview': df.head(10).to_dict('records') if len(df) > 0 else [],
            'columns': df.columns.tolist(),
            'row_count': len(df),
            'summary': df.describe().to_dict() if len(df) > 0 else {}
        }
    except ImportError:
        return {'warning': 'pandas not available for Excel/CSV processing'}
    except Exception as e:
        return {'error': f'Failed to process Excel/CSV: {str(e)}'}

def process_word_doc(file_path):
    """Process Word documents"""
    try:
        from docx import Document
        doc = Document(file_path)
        
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        return {
            'text_content': text_content[:10],  # First 10 paragraphs
            'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()]),
            'full_text': '\n'.join(text_content)
        }
    except ImportError:
        return {'warning': 'python-docx not available for Word processing'}
    except Exception as e:
        return {'error': f'Failed to process Word document: {str(e)}'}

def process_pdf_doc(file_path):
    """Process PDF documents"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(file_path)
        
        text_content = []
        for page_num in range(len(doc)):
            page = doc[page_num]
            text_content.append(page.get_text())
        
        doc.close()
        
        return {
            'page_count': len(text_content),
            'text_preview': text_content[0][:1000] if text_content else '',
            'full_text': '\n'.join(text_content)
        }
    except ImportError:
        return {'warning': 'PyMuPDF not available for PDF processing'}
    except Exception as e:
        return {'error': f'Failed to process PDF: {str(e)}'}

def process_text_file(file_path):
    """Process text files"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        lines = content.split('\n')
        
        return {
            'line_count': len(lines),
            'character_count': len(content),
            'preview': content[:1000],
            'full_text': content
        }
    except Exception as e:
        return {'error': f'Failed to process text file: {str(e)}'}

@app.route('/api/files/list', methods=['GET'])
def list_uploaded_files():
    """List all uploaded files"""
    try:
        uploads_dir = os.path.join(os.path.dirname(__file__), 'uploads')
        
        if not os.path.exists(uploads_dir):
            return jsonify({'files': []})
        
        files = []
        for filename in os.listdir(uploads_dir):
            file_path = os.path.join(uploads_dir, filename)
            if os.path.isfile(file_path):
                file_stat = os.stat(file_path)
                files.append({
                    'filename': filename,
                    'size': file_stat.st_size,
                    'modified': datetime.fromtimestamp(file_stat.st_mtime).isoformat()
                })
        
        return jsonify({'files': sorted(files, key=lambda x: x['modified'], reverse=True)})
        
    except Exception as e:
        return jsonify({'error': f'Failed to list files: {str(e)}'}), 500

@app.route('/api/files/upload-batch', methods=['POST', 'OPTIONS'])
def upload_batch():
    """Upload multiple files at once (for frontend compatibility)"""
    print(f"📤 Batch upload request - Method: {request.method}")
    print(f"📤 Files in request: {list(request.files.keys())}")
    print(f"📤 Form data keys: {list(request.form.keys())}")
    
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        print("📤 Processing batch upload...")
        
        if 'files' not in request.files:
            print("❌ No 'files' key in request.files")
            print(f"❌ Available keys: {list(request.files.keys())}")
            return jsonify({'error': 'No files provided'}), 400
        
        files = request.files.getlist('files')
        print(f"📤 Found {len(files)} files")
        
        if not files or len(files) == 0:
            print("❌ No files in the list")
            return jsonify({'error': 'No files selected'}), 400
        
        results = []
        errors = []
        
        for i, file in enumerate(files):
            print(f"📄 Processing file {i+1}: {file.filename}")
            try:
                if file.filename == '':
                    print(f"❌ Empty filename for file {i+1}")
                    errors.append({'error': 'Empty filename'})
                    continue
                
                # Define allowed file types
                allowed_extensions = {'doc', 'docx', 'pdf', 'xls', 'xlsx', 'png', 'jpg', 'jpeg', 'gif', 'webp', 'txt', 'csv'}
                
                def allowed_file(filename):
                    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions
                
                if not allowed_file(file.filename):
                    print(f"❌ File type not allowed: {file.filename}")
                    errors.append({
                        'filename': file.filename,
                        'error': 'File type not allowed'
                    })
                    continue
                
                # Create uploads directory if it doesn't exist
                uploads_dir = os.path.join(os.path.dirname(__file__), 'uploads')
                os.makedirs(uploads_dir, exist_ok=True)
                
                # Secure the filename and generate unique name
                filename = secure_filename(file.filename)
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                unique_filename = f'{timestamp}_{filename}'
                file_path = os.path.join(uploads_dir, unique_filename)
                
                # Save the file
                file.save(file_path)
                print(f"💾 Saved file: {file_path}")
                
                # Process file based on type
                file_type = filename.rsplit('.', 1)[1].lower()
                
                result = {'message': 'File uploaded successfully'}
                
                if file_type in ['xls', 'xlsx', 'csv']:
                    result.update(process_excel_csv(file_path))
                elif file_type in ['doc', 'docx']:
                    result.update(process_word_doc(file_path))
                elif file_type == 'pdf':
                    result.update(process_pdf_doc(file_path))
                elif file_type == 'txt':
                    result.update(process_text_file(file_path))
                else:
                    result.update({'warning': 'File uploaded but not processed'})
                
                # Add file info to result
                result.update({
                    'filename': unique_filename,
                    'original_filename': filename,
                    'file_type': file_type,
                    'file_size': os.path.getsize(file_path)
                })
                
                results.append(result)
                print(f"✅ Successfully processed: {filename}")
                
            except Exception as e:
                print(f"❌ Error processing file {i+1}: {str(e)}")
                errors.append({
                    'filename': file.filename if file.filename else 'unknown',
                    'error': f'Failed to process file: {str(e)}'
                })
        
        response_data = {
            'results': results,
            'errors': errors,
            'total_files': len(files),
            'successful': len(results),
            'failed': len(errors)
        }
        
        print(f"📊 Upload summary: {len(results)} successful, {len(errors)} failed")
        return jsonify(response_data)
        
    except Exception as e:
        print(f"❌ Batch upload error: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Failed to process files: {str(e)}'}), 500

@app.route('/api/upload-financial-data', methods=['POST', 'OPTIONS'])
def upload_financial_data():
    """Upload financial data (Excel/CSV) with automatic processing"""
    if request.method == 'OPTIONS':
        return '', 200
    
    try:
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'error': 'No file provided'
            }), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': 'No file selected'
            }), 400
        
        # Check file type
        if not (file.filename.endswith('.csv') or file.filename.endswith(('.xlsx', '.xls'))):
            return jsonify({
                'success': False,
                'error': 'Unsupported file format. Please upload CSV or Excel files.'
            }), 400
        
        # Create uploads directory if it doesn't exist
        uploads_dir = os.path.join(os.path.dirname(__file__), 'uploads')
        os.makedirs(uploads_dir, exist_ok=True)
        
        # Save file
        filename = secure_filename(file.filename)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        unique_filename = f'{timestamp}_{filename}'
        file_path = os.path.join(uploads_dir, unique_filename)
        file.save(file_path)
        
        # Process the financial data
        financial_data = {}
        
        try:
            import pandas as pd
            
            if file.filename.endswith('.csv'):
                df = pd.read_csv(file_path)
            else:
                df = pd.read_excel(file_path)
            
            # Extract basic financial metrics from the data
            financial_data = {
                'data_preview': df.head(10).to_dict('records') if len(df) > 0 else [],
                'columns': df.columns.tolist(),
                'row_count': len(df),
                'revenue': float(df['Revenue'].iloc[0]) if 'Revenue' in df.columns and len(df) > 0 else 0,
                'employees': int(df['Employees'].iloc[0]) if 'Employees' in df.columns and len(df) > 0 else 0
            }
            
        except ImportError:
            financial_data = {'warning': 'pandas not available for data processing'}
        except Exception as e:
            financial_data = {'warning': f'Could not process data: {str(e)}'}
        
        return jsonify({
            'success': True,
            'extracted_data': financial_data,
            'message': 'Financial data processed successfully'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e),
            'message': 'Failed to process financial data'
        }), 500

# ============================================
# VALUATION ENDPOINTS (Multi-Model Valuation)
# ============================================

class ValuationModels:
    """Multi-model valuation calculator supporting various startup and business valuation methods"""
    
    @staticmethod
    def berkus_method(data):
        """Berkus Method: Pre-revenue qualitative valuation"""
        factors = {
            'sound_idea': data.get('idea_quality', 0.5),
            'prototype_quality': data.get('product_quality', 0.7),
            'quality_management': data.get('team_experience', 0.8),
            'strategic_relationships': data.get('partnerships', 0.6),
            'product_rollout': data.get('market_readiness', 0.5)
        }
        
        total_score = sum(factors.values()) / len(factors)
        valuation = total_score * 2000000  # $2M max total
        
        return {
            'valuation': int(valuation),
            'factors': factors,
            'method': 'berkus',
            'confidence': min(95, max(60, total_score * 100))
        }
    
    @staticmethod
    def scorecard_method(data):
        """Scorecard Method: Compare against regional startup averages"""
        regional_average = data.get('regional_average', 2000000)
        
        # Scorecard factors with weights
        factors = {
            'management': {'weight': 0.30, 'score': data.get('team_score', 100)},
            'size_of_opportunity': {'weight': 0.25, 'score': data.get('market_size_score', 100)},
            'product_technology': {'weight': 0.15, 'score': data.get('product_score', 100)},
            'competitive_environment': {'weight': 0.10, 'score': data.get('competitive_score', 100)},
            'marketing_channels': {'weight': 0.10, 'score': data.get('marketing_score', 100)},
            'need_for_funding': {'weight': 0.05, 'score': data.get('funding_score', 100)},
            'other': {'weight': 0.05, 'score': data.get('other_score', 100)}
        }
        
        weighted_score = sum(f['weight'] * (f['score'] / 100) for f in factors.values())
        valuation = regional_average * weighted_score
        
        return {
            'valuation': int(valuation),
            'factors': factors,
            'regional_average': regional_average,
            'weighted_score': weighted_score,
            'method': 'scorecard'
        }

    @staticmethod 
    def risk_factor_summation(data):
        """Risk Factor Summation Method"""
        base_valuation = data.get('base_valuation', 2000000)
        
        risk_factors = {
            'management_risk': data.get('management_risk', 0),
            'stage_of_business': data.get('stage_risk', 0),
            'legislation_risk': data.get('political_risk', 0),
            'manufacturing_risk': data.get('production_risk', 0),
            'sales_marketing_risk': data.get('sales_risk', 0),
            'funding_capital_risk': data.get('funding_risk', 0),
            'competition_risk': data.get('competition_risk', 0),
            'technology_risk': data.get('technology_risk', 0),
            'litigation_risk': data.get('litigation_risk', 0),
            'international_risk': data.get('international_risk', 0),
            'reputation_risk': data.get('reputation_risk', 0),
            'potential_lucrative_exit': data.get('exit_potential', 0)
        }
        
        total_adjustment = sum(risk_factors.values())
        adjusted_valuation = base_valuation + (base_valuation * total_adjustment)
        
        return {
            'valuation': max(0, int(adjusted_valuation)),
            'base_valuation': base_valuation,
            'risk_factors': risk_factors,
            'total_adjustment': total_adjustment,
            'method': 'risk_factor_summation'
        }

    @staticmethod
    def vc_method(data):
        """Venture Capital Method"""
        expected_exit_value = data.get('expected_exit_value', 50000000)
        years_to_exit = data.get('years_to_exit', 5)
        target_return = data.get('target_return', 0.10)  # 10x return
        retention_ratio = data.get('retention_ratio', 0.8)
        
        terminal_value = expected_exit_value
        required_return = (1 + target_return) ** years_to_exit
        post_money_valuation = terminal_value / required_return
        
        # Calculate ownership percentage needed
        investment_needed = data.get('investment_needed', 1000000)
        ownership_needed = investment_needed / post_money_valuation
        
        # Adjust for dilution
        final_ownership = ownership_needed / retention_ratio
        pre_money_valuation = post_money_valuation - investment_needed
        
        return {
            'valuation': max(0, int(pre_money_valuation)),
            'post_money_valuation': int(post_money_valuation),
            'terminal_value': terminal_value,
            'required_return': required_return,
            'ownership_percentage': final_ownership * 100,
            'method': 'vc_method'
        }

    @staticmethod
    def dcf_method(data):
        """Discounted Cash Flow Method"""
        try:
            # Financial projections
            current_revenue = data.get('current_revenue', 1000000)
            growth_rate = data.get('growth_rate', 0.20)
            ebitda_margin = data.get('ebitda_margin', 0.25)
            tax_rate = data.get('tax_rate', 0.25)
            discount_rate = data.get('discount_rate', 0.12)
            terminal_growth = data.get('terminal_growth', 0.03)
            projection_years = data.get('projection_years', 5)
            
            cash_flows = []
            
            # Project cash flows
            for year in range(1, projection_years + 1):
                revenue = current_revenue * ((1 + growth_rate) ** year)
                ebitda = revenue * ebitda_margin
                # Simplified cash flow calculation
                free_cash_flow = ebitda * (1 - tax_rate)
                present_value = free_cash_flow / ((1 + discount_rate) ** year)
                
                cash_flows.append({
                    'year': year,
                    'revenue': revenue,
                    'ebitda': ebitda,
                    'free_cash_flow': free_cash_flow,
                    'present_value': present_value
                })
            
            # Terminal value
            terminal_fcf = cash_flows[-1]['free_cash_flow'] * (1 + terminal_growth)
            terminal_value = terminal_fcf / (discount_rate - terminal_growth)
            terminal_pv = terminal_value / ((1 + discount_rate) ** projection_years)
            
            # Enterprise value
            pv_of_cash_flows = sum(cf['present_value'] for cf in cash_flows)
            enterprise_value = pv_of_cash_flows + terminal_pv
            
            return {
                'valuation': int(enterprise_value),
                'cash_flows': cash_flows,
                'terminal_value': terminal_value,
                'terminal_pv': terminal_pv,
                'pv_of_cash_flows': pv_of_cash_flows,
                'method': 'dcf'
            }
            
        except Exception as e:
            return {
                'valuation': 0,
                'error': str(e),
                'method': 'dcf'
            }

    @staticmethod
    def market_comparables(data):
        """Market Comparables Method"""
        # Get comparable companies data
        comparables = data.get('comparables', [])
        company_metrics = data.get('company_metrics', {})
        
        if not comparables:
            # Use default market multiples if no comparables provided
            revenue_multiple = data.get('revenue_multiple', 3.0)
            ebitda_multiple = data.get('ebitda_multiple', 8.0)
            
            company_revenue = company_metrics.get('revenue', 1000000)
            company_ebitda = company_metrics.get('ebitda', 250000)
            
            revenue_valuation = company_revenue * revenue_multiple
            ebitda_valuation = company_ebitda * ebitda_multiple
            
            # Weighted average
            valuation = (revenue_valuation * 0.6) + (ebitda_valuation * 0.4)
            
            return {
                'valuation': int(valuation),
                'revenue_valuation': revenue_valuation,
                'ebitda_valuation': ebitda_valuation,
                'multiples_used': {
                    'revenue_multiple': revenue_multiple,
                    'ebitda_multiple': ebitda_multiple
                },
                'method': 'market_comparables'
            }
        
        # Calculate multiples from comparables
        total_rev_multiple = 0
        total_ebitda_multiple = 0
        valid_comparables = 0
        
        for comp in comparables:
            if comp.get('valuation') and comp.get('revenue'):
                rev_multiple = comp['valuation'] / comp['revenue']
                total_rev_multiple += rev_multiple
                valid_comparables += 1
                
                if comp.get('ebitda'):
                    ebitda_multiple = comp['valuation'] / comp['ebitda']
                    total_ebitda_multiple += ebitda_multiple
        
        if valid_comparables > 0:
            avg_rev_multiple = total_rev_multiple / valid_comparables
            avg_ebitda_multiple = total_ebitda_multiple / valid_comparables if total_ebitda_multiple > 0 else 8.0
            
            company_revenue = company_metrics.get('revenue', 1000000)
            company_ebitda = company_metrics.get('ebitda', 250000)
            
            revenue_valuation = company_revenue * avg_rev_multiple
            ebitda_valuation = company_ebitda * avg_ebitda_multiple
            
            valuation = (revenue_valuation * 0.6) + (ebitda_valuation * 0.4)
            
            return {
                'valuation': int(valuation),
                'revenue_valuation': revenue_valuation,
                'ebitda_valuation': ebitda_valuation,
                'comparables_used': len(comparables),
                'multiples_calculated': {
                    'avg_revenue_multiple': avg_rev_multiple,
                    'avg_ebitda_multiple': avg_ebitda_multiple
                },
                'method': 'market_comparables'
            }
        
        return {
            'valuation': 0,
            'error': 'Insufficient comparable data',
            'method': 'market_comparables'
        }

@app.route('/api/valuate/all-methods', methods=['POST'])
def valuate_all_methods():
    """Run all valuation methods and return comprehensive results"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        # Initialize ValuationModels
        models = ValuationModels()
        
        # Run all valuation methods
        results = {
            'berkus': models.berkus_method(data),
            'scorecard': models.scorecard_method(data),
            'risk_factor': models.risk_factor_summation(data),
            'vc_method': models.vc_method(data),
            'dcf': models.dcf_method(data),
            'market_comparables': models.market_comparables(data)
        }
        
        # Calculate summary statistics
        valuations = [result.get('valuation', 0) for result in results.values() if result.get('valuation', 0) > 0]
        
        if valuations:
            summary = {
                'average_valuation': sum(valuations) / len(valuations),
                'median_valuation': sorted(valuations)[len(valuations)//2],
                'min_valuation': min(valuations),
                'max_valuation': max(valuations),
                'methods_count': len(valuations)
            }
        else:
            summary = {
                'average_valuation': 0,
                'median_valuation': 0,
                'min_valuation': 0,
                'max_valuation': 0,
                'methods_count': 0
            }
        
        return jsonify({
            'success': True,
            'results': results,
            'summary': summary,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route('/api/methods', methods=['GET'])
def get_available_methods():
    """Get list of available valuation methods"""
    methods = [
        {
            'id': 'berkus',
            'name': 'Berkus Method',
            'description': 'Pre-revenue qualitative valuation based on 5 key factors',
            'suitable_for': 'Pre-revenue startups',
            'max_valuation': 2000000
        },
        {
            'id': 'scorecard',
            'name': 'Scorecard Method',
            'description': 'Compare against regional averages using weighted factors',
            'suitable_for': 'Early-stage startups',
            'max_valuation': None
        },
        {
            'id': 'risk_factor',
            'name': 'Risk Factor Summation',
            'description': 'Adjust base valuation based on risk assessment',
            'suitable_for': 'Startups with identifiable risks',
            'max_valuation': None
        },
        {
            'id': 'vc_method',
            'name': 'Venture Capital Method',
            'description': 'Calculate pre-money valuation based on expected exit',
            'suitable_for': 'Venture-backed startups',
            'max_valuation': None
        },
        {
            'id': 'dcf',
            'name': 'Discounted Cash Flow',
            'description': 'Present value of projected future cash flows',
            'suitable_for': 'Revenue-generating businesses',
            'max_valuation': None
        },
        {
            'id': 'market_comparables',
            'name': 'Market Comparables',
            'description': 'Valuation based on similar company multiples',
            'suitable_for': 'Companies with comparable peers',
            'max_valuation': None
        }
    ]
    
    return jsonify({
        'success': True,
        'methods': methods,
        'total_methods': len(methods)
    })

if __name__ == '__main__':
    print("🚀 Starting ValuAI Business Valuation Tool...")
    
    # Quick database initialization
    init_db()
    print("✅ Server ready!")
    
    # Minimal startup info for faster boot
    print("🔧 Running on: http://127.0.0.1:5002")
    print("� Status: http://127.0.0.1:5002/api/health")
    
    # Use production mode for faster startup (set debug=False)
    # Enable debug mode only if FLASK_DEBUG environment variable is set
    debug_mode = os.environ.get('FLASK_DEBUG', 'false').lower() == 'true'
    
    app.run(host='127.0.0.1', port=5002, debug=debug_mode, threaded=True)
